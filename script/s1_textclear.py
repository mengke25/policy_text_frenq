import os
from docx import Document
import pandas as pd
import numpy as np
import jieba
from func_citycode import citycode  # 导入 citycode 函数
from func_provcode import provcode  # 导入 provcode 函数
from func_codetocity import code_to_city  # 导入 code_to_city 函数
from func_codetoprov import code_to_prov  # 导入 code_to_prov 函数

######################################################################################################

# 输入的关键词列表
keywords = ["数据跨境", "数据流动", "数据共享"]

######################################################################################################

# 定义文件路径
base_path = "D:\\py_proj\\policy_analyse"
pool_path = os.path.join(base_path, "pool")
output_path = os.path.join(base_path, "output", "policy_list.xlsx")
stopword_path = os.path.join(base_path, "files", "stopword.txt")
custom_dict_path = os.path.join(base_path, "files", "custom.txt")

# 加载自定义词典和停用词
jieba.load_userdict(custom_dict_path)

with open(stopword_path, 'r', encoding='utf-8') as f:
    stopwords = set(f.read().strip().split())


################################################################################################
def clean_text(text):
    # 移除无效字符
    return ''.join(char for char in text if char.isprintable() and char not in {'\x00', '\x01', '\x02', '\x03', '\x04', '\x05', '\x06', '\x07', '\x08', '\x0b', '\x0c', '\x0e', '\x0f', '\x10', '\x11', '\x12', '\x13', '\x14', '\x15', '\x16', '\x17', '\x18', '\x19', '\x1a', '\x1b', '\x1c', '\x1d', '\x1e', '\x1f'})

def extract_year_from_text(text):
    for i in range(len(text) - 3):
        year = text[i:i+4]
        if year.isdigit() and 2010 <= int(year) <= 2025:
            return year
    return None

def extract_year_from_filename(filename):
    year = extract_year_from_text(filename)
    return year


################################################################################################
# 步骤一：将txt文件转换为docx文件
txt_files = []
for root, dirs, files in os.walk(pool_path):
    for file in files:
        if file.endswith('.txt'):
            txt_files.append(os.path.join(root, file))

for i, txt_path in enumerate(txt_files, start=1):
    docx_path = txt_path.replace('.txt', '.docx')

    if not os.path.exists(docx_path):
        with open(txt_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # 清理内容
        content = clean_text(content)

        doc = Document()
        doc.add_paragraph(content)
        doc.save(docx_path)
        print(f"正在转化txt文件 {i}/{len(txt_files)} 已完成 {i/len(txt_files)*100:.2f}%")
    else:
        print(f"跳过已存在的docx文件: {docx_path}")

# 打印输入的关键词信息
print(f"共输入{len(keywords)}个关键词，分别是{keywords}，开始提取")

################################################################################################
# 步骤二：分析docx文件并生成Excel文件
data = []
docx_files = []
for root, dirs, files in os.walk(pool_path):
    for file in files:
        if file.endswith('.docx'):
            docx_files.append(os.path.join(root, file))

for i, docx_path in enumerate(docx_files, start=1):
    doc = Document(docx_path)
    content = "\n".join([para.text for para in doc.paragraphs])

    # 从文件名提取年份
    year = extract_year_from_filename(os.path.basename(docx_path))
    if not year:
        # 从内容中提取年份
        year = extract_year_from_text(content)

    # 计算总词数（排除停用词）
    words = jieba.lcut(content)
    words_filtered = [word for word in words if word not in stopwords]
    total_word_count = len(words_filtered)

    # 计算关键词出现频次
    keyword_counts = {kw: content.count(kw) for kw in keywords}

    # 获取文件的相对路径（相对于 pool_path）
    relative_path = os.path.relpath(docx_path, pool_path)
    folder_name = os.path.dirname(relative_path)  # 获取文件所在的子文件夹

    # 添加到数据列表
    data.append([
        os.path.basename(docx_path),
        year,
        total_word_count,
        *keyword_counts.values(),
        folder_name
    ])
    print(f"正在处理docx文件 {i}/{len(docx_files)} 已完成 {i/len(docx_files)*100:.2f}%")

# 创建DataFrame
columns = ["文件名", "年份", "总词数"] + keywords + ["文件来源"]
df = pd.DataFrame(data, columns=columns)

################################################################################################
# 调用citycode、provcode函数进行城市、省份代码的提取
df = citycode(df, '文件名')
df.rename(columns={'citycode': 'citycode1'}, inplace=True)
df = provcode(df, '文件名')
df.rename(columns={'provcode': 'provcode1'}, inplace=True)
df = provcode(df, '文件来源')
df.rename(columns={'provcode': 'provcode2'}, inplace=True)

################################################################################################
# 数据处理
# Replace citycode1 and provcode1 with NaN if they are 0
df['citycode1'].replace(0, np.nan, inplace=True)
df['provcode1'].replace(0, np.nan, inplace=True)
df['provcode2'].replace(0, np.nan, inplace=True)

# Replace provcode1 with provcode2 if provcode1 is NaN and provcode2 is not NaN
if 'provcode2' in df.columns:
    df['provcode1'] = df.apply(lambda row: row['provcode2'] if pd.isna(row['provcode1']) and not pd.isna(row['provcode2']) else row['provcode1'], axis=1)
    # Drop provcode2
    df.drop(columns=['provcode2'], inplace=True)

# Rename columns
df.rename(columns={'provcode1': 'provcode', 'citycode1': 'citycode'}, inplace=True)

# 将 citycode 的前两位转换为省级代码，并更新 provcode
df['provcode'] = df.apply(lambda row: int(str(row['citycode'])[:2]) * 10000 if not pd.isna(row['citycode']) else row['provcode'], axis=1)

# Generate 政策级别 based on citycode and provcode
df['政策级别'] = df.apply(lambda row: "市级政策" if not pd.isna(row['citycode']) else 
                                      ("省级政策" if pd.isna(row['citycode']) and not pd.isna(row['provcode']) else "未匹配"), axis=1)
df.loc[df['provcode'].isin([310000, 110000, 120000, 500000]), '政策级别'] = "直辖市政策"


df = code_to_city(df, 'citycode')
df = code_to_prov(df, 'provcode')

################################################################################################
# 保存到Excel文件
os.makedirs(os.path.dirname(output_path), exist_ok=True)
df.to_excel(output_path, index=False)

print(f"分析结果已保存到 {output_path}")
